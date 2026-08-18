#!/usr/bin/env python3
"""Builds the editable Word version of a WAGOLL wall.

    pip install python-docx
    python3 tools/build_docx.py pigeon-patrol
    python3 tools/build_docx.py gs73 batman

The wall text is read straight out of that topic's content.js, the same file
the web pages and the PDFs use, so the Word document can never drift from
them. Nothing here holds a copy of the wording.

Each argument is a key into SUBJECTS below, not a folder path: the topics live
several directories deep (history/gs73, history/batman, economics/pigeon-patrol)
and every output lands in print/ alongside the PDFs, so neither end of the job
is "the subject's own folder" any more. Add a topic here when it gets a wall.

The layout deliberately mirrors the original Year 7 Economics WAGOLL .docx:
landscape, one table, criterion colours applied to runs, so it opens looking
like the file it replaces and can be edited in Word as usual.
"""

import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm

ROOT = Path(__file__).resolve().parents[1]

# Every topic that has a wall. "content" is relative to the repo root; every
# output is written to print/, alongside the PDFs build-pdf.mjs makes for the
# same topic.
SUBJECTS = {
    "gs73":          {"content": "history/gs73/content.js",
                       "out": "GS73-WAGOLL-Wall.docx"},
    "batman":        {"content": "history/batman/content.js",
                       "out": "Batman-Treaty-WAGOLL-Wall.docx"},
    "pigeon-patrol": {"content": "economics/pigeon-patrol/content.js",
                       "out": "Pigeon-Patrol-WAGOLL-Wall.docx"},
}

# Criterion colours, matched to the web pages. These are the print/light values.
INK = RGBColor(0x20, 0x23, 0x1F)
COLOURS = {
    "success":    "1F4ED8", "innovation": "C2410C", "decisions": "15803D",
    "source":     "176B87", "context":    "357A46", "evidence":  "A34D1D",
    "judge":      "744F91", "meta":       "7C3AED",
}
# gs73's metacognition is the grey-blue, pigeon-patrol's is the violet above
META_BY_SUBJECT = {"gs73": "53666B", "pigeon-patrol": "7C3AED"}
FILLS = {
    "success":   "E4EBFE", "innovation": "FCEEE6", "decisions": "E2F4E8",
    "source":    "D8EDF4", "context":    "DFF0DF", "evidence":  "F7E3D4",
    "judge":     "EADFF2", "meta":       "E2EAEB",
}
DEEP = "223F43"

# Second channel, same as the web pages: the underline style carries the
# criterion when the colour is gone (photocopied, or colour-blind readers).
UNDERLINES = {
    "success": WD_UNDERLINE.SINGLE, "source":  WD_UNDERLINE.SINGLE,
    "innovation": WD_UNDERLINE.DASH, "context": WD_UNDERLINE.DASH,
    "decisions": WD_UNDERLINE.DOTTED, "evidence": WD_UNDERLINE.DOTTED,
    "judge": WD_UNDERLINE.DOUBLE,
    "meta": WD_UNDERLINE.WAVY,
}


def read_content(content_path: str) -> dict:
    """Evaluate content.js in node and hand the data back as JSON."""
    js = ROOT / content_path
    if not js.exists():
        sys.exit(f"no content.js at {js}")
    script = (
        f"const fs=require('fs'),vm=require('vm');const c={{}};vm.createContext(c);"
        f"vm.runInContext(fs.readFileSync({json.dumps(str(js))},'utf8')"
        f"+';globalThis.OUT=JSON.stringify({{WALL,CRITERIA,LEVELS,EXAMPLES,EXPLANATIONS,CONTINUUM,"
        f"GAP:(typeof GAP_TEXT!==\"undefined\")?GAP_TEXT:\"\"}});',c);"
        f"process.stdout.write(c.OUT);"
    )
    out = subprocess.run(["node", "-e", script], capture_output=True, text=True)
    if out.returncode:
        sys.exit("could not read content.js:\n" + out.stderr)
    return json.loads(out.stdout)


def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def set_col_widths(table, widths_mm):
    """Word needs the width on every cell, not just the column."""
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_mm):
            cell.width = Mm(w)


def write_marked(par, text, colours, size, base=INK):
    """Render {key|phrase} as a coloured, underlined run; everything else stays
    black. The black text is the glue, which is the whole point of the wall."""
    for chunk in re.split(r"(\{\w+\|[^}]*\})", text):
        if not chunk:
            continue
        m = re.fullmatch(r"\{(\w+)\|([^}]*)\}", chunk)
        run = par.add_run(m.group(2) if m else chunk)
        run.font.size = Pt(size)
        if m and m.group(1) in colours:
            run.font.color.rgb = RGBColor.from_string(colours[m.group(1)])
            run.bold = True
            run.font.underline = UNDERLINES.get(m.group(1), WD_UNDERLINE.SINGLE)
        else:
            run.font.color.rgb = base


def plain(text):
    return re.sub(r"\{(\w+)\|([^}]*)\}", r"\2", text)


def build(subject: str) -> Path:
    if subject not in SUBJECTS:
        sys.exit(f"unknown subject '{subject}': add it to SUBJECTS in this file. "
                  f"Known: {', '.join(sorted(SUBJECTS))}")
    subj = SUBJECTS[subject]
    d = read_content(subj["content"])
    wall, crits, levels = d["WALL"], d["CRITERIA"], d["LEVELS"]
    colours = dict(COLOURS)
    colours["meta"] = META_BY_SUBJECT.get(subject, COLOURS["meta"])

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Mm(420), Mm(297)          # A3 landscape
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, side, Mm(10))
    usable = 420 - 20

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(8)
    style.paragraph_format.space_after = Pt(0)

    title = doc.add_paragraph()
    r = title.add_run(re.sub("<[^>]+>", "", wall["title"]))
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(DEEP)

    sub = doc.add_paragraph()
    r = sub.add_run(re.sub("<[^>]+>", "", wall.get("task") or wall.get("inquiry") or ""))
    r.font.size = Pt(9)
    sub.paragraph_format.space_after = Pt(6)

    ncols = len(levels) + 1
    label_w = 34
    col_w = (usable - label_w) / len(levels)
    rows = 3 + len(crits)
    table = doc.add_table(rows=rows, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    hdr = table.rows[0]
    for i, text in enumerate([""] + levels):
        cell = hdr.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, "A16207" if text == wall.get("expected") else DEEP)

    def row_label(cell, text, colour=None):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(colour) if colour else RGBColor.from_string(DEEP)

    # worked examples
    row_label(table.rows[1].cells[0], "WORKED EXAMPLE")
    for i, lv in enumerate(levels, start=1):
        write_marked(table.rows[1].cells[i].paragraphs[0], d["EXAMPLES"][lv], colours, 7.5)

    # why it meets the level
    row_label(table.rows[2].cells[0], "WHAT IT PROVES")
    for i, lv in enumerate(levels, start=1):
        cell = table.rows[2].cells[i]
        shade(cell, "F5F0E6")
        entry = d["EXPLANATIONS"][lv]
        items = entry.get("items", entry) if isinstance(entry, dict) else entry
        first = True
        for c in crits:
            v = items.get(c["key"])
            if v is None:
                continue
            label, text = (v[0], v[1]) if isinstance(v, list) else ("", v)
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            g = p.add_run(c["glyph"] + " ")
            g.font.size = Pt(7)
            g.font.color.rgb = RGBColor.from_string(colours[c["key"]])
            if label:
                b = p.add_run(label + ": ")
                b.bold = True
                b.font.size = Pt(7)
                b.font.color.rgb = RGBColor.from_string(colours[c["key"]])
            t = p.add_run(text)
            t.font.size = Pt(7)
            t.font.color.rgb = INK

    # continuum
    for j, c in enumerate(crits):
        row = table.rows[3 + j]
        row_label(row.cells[0], re.sub("&amp;", "&", c["row"]).upper(), colours[c["key"]])
        for i, lv in enumerate(levels, start=1):
            cell = row.cells[i]
            shade(cell, FILLS[c["key"]])
            text = d["CONTINUUM"][c["key"]].get(lv) or d.get("GAP") or ":"
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(7.5)
            run.font.color.rgb = INK
            if not d["CONTINUUM"][c["key"]].get(lv):
                run.italic = True

    set_col_widths(table, [label_w] + [col_w] * len(levels))

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(6)
    fr = foot.add_run(re.sub("<[^>]+>", "", wall["foot"]))
    fr.font.size = Pt(6.5)
    fr.font.color.rgb = RGBColor(0x67, 0x5F, 0x54)

    out = ROOT / "print" / subj["out"]
    doc.save(out)
    return out


if __name__ == "__main__":
    targets = sys.argv[1:] or list(SUBJECTS)
    for t in targets:
        print("wrote", build(t).relative_to(ROOT))
